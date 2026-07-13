from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn

_CAPTION_MAX = 90


_ROLES = {
    "head":              ("head", "ФИО и должность руководителя", True),
    "head_fio":          ("head_fio", "ФИО руководителя", True),
    "head_position":     ("head_position", "Должность руководителя", False),
    "employee_fio":      ("employee_fio", "ФИО работника", True),
    "employee_position": ("employee_position", "Должность работника", False),
    "department":        ("department", "Подразделение", True),
    "reason":            ("reason", "Причина / основание", True),
    "field":             ("field", "Поле", False),
}


def _norm(s: str) -> str:
    return re.sub(r"\s+", " ", s or "").strip().lower()


def _looks_like_caption(t: str) -> bool:
    t = (t or "").strip()
    if not t:
        return False
    if t.startswith("("):
        return True
    if len(t) > _CAPTION_MAX:
        return False
    return (
        t[:1].islower()
        or bool(re.search(r"ф\.?и\.?о|должност|подпис|подразделен|наименован|причин|руководител|работник", t, re.I))
    )


def _caption_entity(text: str) -> str | None:
    """Сущность подписи для СМЕЖНОСТИ (шапка ФИО/должность)."""
    low = (text or "").lower()
    if "подпис" in low:
        return None
    if re.search(r"руководител|ректор|директор|проректор", low):
        return "head"
    if re.search(r"работник|сотрудник", low):
        return "employee"
    if re.search(r"подразделен|наименован", low):
        return "department"
    return None


def _caption_role(caption: str, prec: str = "") -> str | None:
    """Роль для поля, у которого подпись/контекст рядом (для абзацев и подчёркнутых
    пробелов, где нет пары ФИО-над/должность-под)."""
    cap = f"{prec} {caption}".lower()
    if re.search(r"сведени|второ\w*\s+родител|подпис", cap):
        return None
    has_fio = bool(re.search(r"ф\.?\s*и\.?\s*о|фамили", cap))
    has_pos = "должност" in cap
    if re.search(r"\bкому\b|руководител|ректор|директор|проректор", cap):
        return "head"
    if re.search(r"работник|сотрудник", cap):
        return "employee_fio" if has_fio or not has_pos else "employee_position"
    if re.search(r"подразделен", cap) or (re.search(r"наименован", cap)):
        return "department"
    if re.search(r"причин|основани|уволить", cap):
        return "reason"
    if has_pos:
        return "employee_position"
    if has_fio:
        return "employee_fio"
    return None


# ---------------------------------------------------------------------------
# Детект областей ввода
# ---------------------------------------------------------------------------
def _is_field_table(tbl) -> bool:
    """1×1 таблица — область ввода (в эталонах — с нижней рамкой)."""
    return len(tbl.rows) == 1 and len(tbl.columns) == 1


def _para_has_border(p) -> bool:
    pPr = p._p.find(qn("w:pPr"))
    if pPr is None:
        return False
    bdr = pPr.find(qn("w:pBdr"))
    return bdr is not None and (bdr.find(qn("w:top")) is not None or bdr.find(qn("w:bottom")) is not None)


def _is_border_empty_para(paras, i) -> bool:
    """Пустой абзац-линия: сам с рамкой ИЛИ следующий рисует линию сверху (гос-формы)."""
    p = paras[i]
    if p.text.strip():
        return False
    if _para_has_border(p):
        return True
    return i + 1 < len(paras) and _para_has_border(paras[i + 1])


def _uspace_runs(paragraph):
    """Индексы runs абзаца, являющихся ПОДЧЁРКНУТЫМИ ПРОБЕЛАМИ (место ввода)."""
    out = []
    for k, r in enumerate(paragraph.runs):
        txt = r.text or ""
        if r.underline and txt and set(txt) <= {" ", "\t", "\xa0"} and len(txt) >= 2:
            out.append(k)
    return out


def _ordered_body(doc):
    """Тело документа в порядке: ('p', paragraph) | ('tbl', table)."""
    p_iter, t_iter = iter(doc.paragraphs), iter(doc.tables)
    out = []
    for child in doc.element.body:
        if child.tag == qn("w:p"):
            out.append(("p", next(p_iter)))
        elif child.tag == qn("w:tbl"):
            out.append(("tbl", next(t_iter)))
    return out


# ---------------------------------------------------------------------------
# Анализ: план заполнения + пользовательская схема
# ---------------------------------------------------------------------------
def analyze(path: str | Path) -> tuple[list[dict], list[dict]]:
    """Возвращает (schema, plan).
    plan[i] = {kind:'table'|'para'|'uspace', ref..., field}
    schema — пользовательские поля {name,label,type,required}."""
    doc = Document(str(path))
    ordered = _ordered_body(doc)
    paras = doc.paragraphs

    schema: dict[str, dict] = {}
    role_counts: dict[str, int] = {}

    def field_for(role_key: str) -> str:
        base, label, req = _ROLES.get(role_key, _ROLES["field"])
        role_counts[base] = role_counts.get(base, 0) + 1
        name = base if role_counts[base] == 1 else f"{base}_{role_counts[base]}"
        if name not in schema:
            lbl = label if role_counts[base] == 1 else f"{label} {role_counts[base]}"
            schema[name] = {"name": name, "label": lbl, "type": "string", "required": req}
        return name

    plan: list[dict] = []

    # === 1) ТАБЛИЦЫ (1×1) — роли по СМЕЖНОСТИ подписи (ФИО над / должность под) ===
    # seq: последовательность значимых слотов — таблицы-поля (с ИНДЕКСОМ в doc.tables)
    # и подписи-с-сущностью.
    seq = []  # (kind:'tbl'|'cap', table_idx|entity)
    tbl_counter = 0
    for kind, ref in ordered:
        if kind == "tbl":
            if _is_field_table(ref):
                seq.append(("tbl", tbl_counter))
            tbl_counter += 1
        elif kind == "p":
            ts = ref.text.strip()
            if _looks_like_caption(ts):
                ent = _caption_entity(ts)
                if ent:
                    seq.append(("cap", ent))
    tbl_role: dict[int, str] = {}   # table_idx → role
    for ci, s in enumerate(seq):
        if s[0] != "cap":
            continue
        ent = s[1]
        above = seq[ci - 1] if ci - 1 >= 0 and seq[ci - 1][0] == "tbl" else None
        below = seq[ci + 1] if ci + 1 < len(seq) and seq[ci + 1][0] == "tbl" else None
        if ent in ("head", "employee"):
            if above and above[1] not in tbl_role:
                tbl_role[above[1]] = f"{ent}_fio"
            if below and below[1] not in tbl_role:
                tbl_role[below[1]] = f"{ent}_position"
        else:  # department
            if above and above[1] not in tbl_role:
                tbl_role[above[1]] = "department"
            if below and below[1] not in tbl_role:
                tbl_role[below[1]] = "department"

    for kind, tidx in seq:
        if kind == "tbl" and tidx in tbl_role:
            plan.append({"kind": "table", "table_idx": tidx, "field": field_for(tbl_role[tidx])})

    # === 2) ВЫДЕЛЕННЫЕ ПУСТЫЕ АБЗАЦЫ (гос-формы) — по ближайшей подписи ниже.
    # Несколько пустых строк под одной подписью (перенос длинного значения) = ОДНО
    # поле (заполняем первую). ===
    def _cap_below(i: int) -> str:
        for j in range(i + 1, min(i + 12, len(paras))):
            t = paras[j].text.strip()
            if not t:
                continue
            return t if _looks_like_caption(t) else ""
        return ""

    epara = [(i, _cap_below(i)) for i in range(len(paras)) if _is_border_empty_para(paras, i)]
    used_para: set[int] = set()
    gi = 0
    while gi < len(epara):
        cap = epara[gi][1]
        grp = [epara[gi]]
        gj = gi + 1
        while gj < len(epara) and epara[gj][1] == cap:
            grp.append(epara[gj])
            gj += 1
        role = _caption_role(cap)
        if role:
            plan.append({"kind": "para", "para_idx": grp[0][0], "field": field_for(role)})
            used_para.update(p[0] for p in grp)
        gi = gj

    # === 3) ПОДЧЁРКНУТЫЕ ПРОБЕЛЫ — по тексту перед ними / подписи ниже ===
    used_uspace_para = {p["para_idx"] for p in plan if p["kind"] == "para"}
    for i, p in enumerate(paras):
        runs_idx = _uspace_runs(p)
        if not runs_idx or i in used_uspace_para:
            continue
        # подпись ниже
        cap = ""
        for j in range(i + 1, min(i + 4, len(paras))):
            t = paras[j].text.strip()
            if not t:
                continue
            cap = t if _looks_like_caption(t) else ""
            break
        # текст перед первым подчёркнутым пробелом
        prec = ""
        acc = 0
        for k, r in enumerate(p.runs):
            if k == runs_idx[0]:
                break
            acc += len(r.text or "")
        prec = " ".join(p.text[:acc].split()[-6:])
        role = _caption_role(cap, prec)
        # только ПЕРВЫЙ подчёркнутый пробел абзаца заполняем значением поля
        if role:
            plan.append({"kind": "uspace", "para_idx": i, "run_idx": runs_idx[0],
                         "field": field_for(role)})

    return list(schema.values()), plan


# ---------------------------------------------------------------------------
# Заполнение
# ---------------------------------------------------------------------------
def _set_cell_value(tbl, value: str) -> None:
    """Значение в ячейку 1×1 по центру (сохраняя рамку/стиль)."""
    cell = tbl.rows[0].cells[0]
    para = cell.paragraphs[0]
    # выравнивание по центру
    pPr = para._p.get_or_add_pPr()
    if pPr.find(qn("w:jc")) is None:
        from docx.oxml import OxmlElement
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "center")
        pPr.append(jc)
    # очистить существующий текст и вписать значение
    for r in list(para.runs):
        r.text = ""
    if para.runs:
        para.runs[0].text = value
    else:
        para.add_run(value)


def _fill_uspace(paragraph, run_idx: int, value: str) -> None:
    """Значение поверх подчёркнутого-пробельного run (подчёркивание сохраняется)."""
    run = paragraph.runs[run_idx]
    width = len(run.text or "")
    v = " ".join(str(value).split())
    # по центру в пределах исходной ширины (не короче значения+2)
    if len(v) + 2 >= width:
        run.text = f" {v} "
    else:
        rem = width - len(v)
        run.text = " " * (rem // 2) + v + " " * (rem - rem // 2)


def autofill(path: str | Path, values: dict[str, Any], out_path: str | Path) -> Path:
    """Вписывает значения в области ввода (таблицы/абзацы/подчёркнутые пробелы)."""
    doc = Document(str(path))
    _, plan = analyze(path)
    paras = doc.paragraphs
    # таблицы в порядке — чтобы сопоставить по индексу с планом
    used: set[str] = set()

    for it in plan:
        fld = it["field"]
        val = values.get(fld)
        if it["kind"] == "table":
            if fld in used:
                continue
            used.add(fld)
            _set_cell_value(doc.tables[it["table_idx"]], str(val) if val not in (None, "") else "")
        elif it["kind"] == "para":
            if val in (None, "") or fld in used:
                continue
            used.add(fld)
            run = paras[it["para_idx"]].add_run(str(val))
            run.underline = True
        elif it["kind"] == "uspace":
            if val in (None, "") or fld in used:
                continue
            used.add(fld)
            _fill_uspace(paras[it["para_idx"]], it["run_idx"], str(val))

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def render_field_preview(src: str | Path, out_path: str | Path) -> Path:
    """Версия бланка ДЛЯ ПОКАЗА: области ввода заполнены названиями полей [в скобках]."""
    schema, plan = analyze(src)
    labels = {f["name"]: f["label"] for f in schema}
    doc = Document(str(src))
    paras = doc.paragraphs
    seen: set[str] = set()
    for it in plan:
        fld = it["field"]
        if fld in seen and it["kind"] != "table":
            continue
        seen.add(fld)
        label = f"[{labels.get(fld, fld)}]"
        if it["kind"] == "table":
            _set_cell_value(doc.tables[it["table_idx"]], label)
        elif it["kind"] == "para":
            run = paras[it["para_idx"]].add_run(label)
            run.underline = True
        elif it["kind"] == "uspace":
            _fill_uspace(paras[it["para_idx"]], it["run_idx"], label)
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def has_jinja_placeholders(path: str | Path) -> bool:
    """True, если в .docx есть {{ переменные }} (обычный docxtpl-путь)."""
    try:
        from docxtpl import DocxTemplate
        return bool(DocxTemplate(str(path)).get_undeclared_template_variables())
    except Exception:
        return False
