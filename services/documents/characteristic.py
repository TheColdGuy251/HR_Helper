"""Б1 (обязательный процесс из раздела «ИИ чат-бот»): характеристика для
представления работника к внешней награде — из ходатайства 1С:Документооборот.

Вход: .docx печатной формы «ХОДАТАЙСТВО О НАГРАЖДЕНИИ» (награда, основание,
ФИО, должность/подразделение, степень/звание, трудовая деятельность, награды
за 5 лет, «Конкретные результаты работы и основные достижения»).
Выход: .docx характеристики в стиле образцов ТИУ (два семейства: ППС и
АУП/специалисты), запись в «Мои документы».

ПДн: текст ходатайства обрабатывается эфемерно и НЕ попадает в базу знаний.
"""
from __future__ import annotations

import json
import re
from datetime import datetime
from pathlib import Path
from typing import Any

from sqlalchemy.orm import Session

from config import settings
from data.my_documents import MyDocuments
from data.users import User
from services.llm import get_llm
from services.llm.prompts import SYSTEM_PROMPT_CHARACTERISTIC, SYSTEM_PROMPT_PETITION_EXTRACT
from utils.logger import logger

# Роли ППС — для них свой стиль характеристики (наука/преподавание).
_PPS_RE = re.compile(
    r"профессор|доцент|преподават|ассистент|заведующ\w*\s+кафедр|научн\w*\s+сотрудник",
    re.IGNORECASE,
)

# Запрос «сделай характеристику» в чате (с учётом склонений).
CHARACTERISTIC_REQUEST_RE = re.compile(
    r"характеристик\w*", re.IGNORECASE
)

_FIELD_KEYS = (
    "award", "basis", "fio", "position", "department",
    "degree", "rank", "career", "awards", "achievements",
)

# Regex-подстраховка по стабильной печатной форме (на случай недоступной LLM).
_AWARD_RE = re.compile(r"Награда\s*[:：]?\s*(.+)", re.IGNORECASE)
_BASIS_RE = re.compile(r"Основание\s*[:：]?\s*(.+)", re.IGNORECASE)
_FIO_ROW_RE = re.compile(
    r"Фамилия,?\s*имя,?\s*отчество\s*[:：|]?\s*([А-ЯЁ][а-яё]+(?:\s+[А-ЯЁ][а-яё]+){1,2})"
)


def detect_category(position: str | None) -> str:
    """'pps' — преподавательские должности, 'aup' — административные/специалисты."""
    return "pps" if position and _PPS_RE.search(position) else "aup"


def _quick_parse(text: str) -> dict[str, Any]:
    out: dict[str, Any] = {}
    for key, rx in (("award", _AWARD_RE), ("basis", _BASIS_RE), ("fio", _FIO_ROW_RE)):
        m = rx.search(text)
        if m:
            out[key] = m.group(1).strip().strip("|").strip()
    return out


def parse_petition(text: str) -> dict[str, Any]:
    """Извлекает поля ходатайства: LLM (основной путь) + regex-подстраховка.
    Любое поле может быть None/[] — UI даёт пользователю поправить руками."""
    fields: dict[str, Any] = {k: None for k in _FIELD_KEYS}
    fields["career"], fields["awards"] = [], []

    sample = (text or "")[:8000]
    try:
        data = get_llm().generate_json(
            SYSTEM_PROMPT_PETITION_EXTRACT, sample,
            schema_hint='{"award": "...", "basis": "...", "fio": "...", "position": "...", '
                        '"department": "...", "degree": null, "rank": null, '
                        '"career": [], "awards": [], "achievements": "..."}',
        )
        if isinstance(data, dict) and not data.get("_mock"):
            for k in _FIELD_KEYS:
                v = data.get(k)
                if k in ("career", "awards"):
                    fields[k] = [str(x).strip() for x in v if str(x).strip()] if isinstance(v, list) else []
                elif isinstance(v, str) and v.strip() and v.strip().lower() != "null":
                    fields[k] = v.strip()
    except Exception as e:
        logger.warning("[CHAR] LLM-извлечение ходатайства не удалось: {}", e)

    for k, v in _quick_parse(sample).items():
        fields[k] = fields[k] or v
    fields["category"] = detect_category(fields.get("position"))
    return fields


def _fields_block(fields: dict[str, Any]) -> str:
    """Данные ходатайства в читаемом виде для промпта генерации."""
    lines = ["Данные из ходатайства о награждении:"]
    labels = {
        "award": "Награда", "basis": "Основание", "fio": "ФИО",
        "position": "Должность", "department": "Подразделение",
        "degree": "Учёная степень", "rank": "Учёное звание",
    }
    for k, label in labels.items():
        if fields.get(k):
            lines.append(f"{label}: {fields[k]}")
    if fields.get("career"):
        lines.append("Трудовая деятельность в ТИУ:")
        lines += [f"- {c}" for c in fields["career"]]
    if fields.get("awards"):
        lines.append("Награды и поощрения за последние 5 лет:")
        lines += [f"- {a}" for a in fields["awards"]]
    if fields.get("achievements"):
        lines.append(f"Конкретные результаты работы и основные достижения:\n{fields['achievements']}")
    return "\n".join(lines)


_STYLE_HINTS = {
    "pps": (
        "Категория работника: профессорско-преподавательский состав. Акценты стиля: научная и "
        "научно-педагогическая работа, публикации/монографии (если указаны), подготовка "
        "обучающихся, участие в советах/комиссиях, вклад в развитие направления."
    ),
    "aup": (
        "Категория работника: административно-управленческий персонал / специалист. Акценты "
        "стиля: карьерный путь («прошёл(шла) путь от … до …»), организация и результаты работы "
        "подразделения, участие в проектах/проверках/аккредитациях, деловые качества."
    ),
}


def generate_characteristic_text(fields: dict[str, Any], category: str | None = None) -> str:
    """Связный текст характеристики по данным ходатайства (LLM, без выдумывания)."""
    cat = category or fields.get("category") or detect_category(fields.get("position"))
    user_msg = _STYLE_HINTS.get(cat, _STYLE_HINTS["aup"]) + "\n\n" + _fields_block(fields)
    llm = get_llm()
    text = llm.generate_text(
        SYSTEM_PROMPT_CHARACTERISTIC, user_msg, max_tokens=1400, temperature=0.3
    )
    text = (text or "").strip()
    if not text:
        raise RuntimeError("LLM не вернула текст характеристики")
    return text


def render_characteristic_docx(fields: dict[str, Any], body_text: str) -> Path:
    """Собирает .docx: шапка «ХАРАКТЕРИСТИКА» + ФИО/должность + абзацы текста."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.add_run("ХАРАКТЕРИСТИКА")
    run.bold = True
    run.font.size = Pt(14)

    sub_bits = [b for b in (fields.get("fio"), fields.get("position"), fields.get("department")) if b]
    if sub_bits:
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = sub.add_run(", ".join(sub_bits))
        r.bold = True
    if fields.get("award"):
        aw = doc.add_paragraph()
        aw.alignment = WD_ALIGN_PARAGRAPH.CENTER
        aw.add_run(f"(для представления к награде: {fields['award']})").italic = True

    doc.add_paragraph()
    for para in re.split(r"\n\s*\n+", body_text):
        p = doc.add_paragraph(para.strip())
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Pt(28)

    settings.docs_generated.mkdir(parents=True, exist_ok=True)
    fio_slug = re.sub(r"[^\w]+", "_", (fields.get("fio") or "работник"))[:40]
    out = settings.docs_generated / f"characteristic_{fio_slug}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(str(out))
    return out


def create_characteristic(
    db: Session, user: User, fields: dict[str, Any], category: str | None = None
) -> tuple[MyDocuments, str]:
    """Полный цикл: текст → docx → запись в «Мои документы». Возвращает (doc, текст)."""
    text = generate_characteristic_text(fields, category)
    path = render_characteristic_docx(fields, text)
    title = "Характеристика"
    if fields.get("fio"):
        title += f" — {fields['fio']}"
    rec = MyDocuments(
        user_id=user.id,
        title=title,
        template_key="characteristic",
        file_path=str(path),
        progress=100,
        status="ready",
        fields={k: fields.get(k) for k in (*_FIELD_KEYS, "category")},
    )
    db.add(rec)
    db.commit()
    db.refresh(rec)
    logger.info("[CHAR] характеристика создана: doc={} ({})", rec.id, title)
    return rec, text


def petition_fields_from_json(raw: str | dict) -> dict[str, Any]:
    """Нормализует поля из запроса клиента (модалка может прислать правки)."""
    data = raw if isinstance(raw, dict) else json.loads(raw or "{}")
    fields: dict[str, Any] = {k: None for k in _FIELD_KEYS}
    fields["career"], fields["awards"] = [], []
    for k in _FIELD_KEYS:
        v = data.get(k)
        if k in ("career", "awards"):
            if isinstance(v, list):
                fields[k] = [str(x).strip() for x in v if str(x).strip()]
            elif isinstance(v, str):
                fields[k] = [s.strip() for s in v.splitlines() if s.strip()]
        elif isinstance(v, str) and v.strip():
            fields[k] = v.strip()
    fields["category"] = data.get("category") or detect_category(fields.get("position"))
    return fields
