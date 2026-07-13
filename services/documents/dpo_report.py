"""Б2 («самое массивное» из раздела «ИИ чат-бот»): отчёт по ДПО из выгрузки 1С:ЗиК.

Вход: xlsx-выгрузка «ПК за год» (шапка с параметрами, затем таблица: Физическое
лицо, Должность, Подразделение, Иерархия, Категория должности, Учреждение,
Вид образования, Группа, даты, Номер документа, Примечание, Тема, Часов, …).
ВАЖНО: 1С группирует строки по человеку — ФИО заполнено только в ПЕРВОЙ строке
блока, дальше forward-fill.

Выход: word-отчёт по структуре образца «ДПО за 2023»: все числа считаются
ДЕТЕРМИНИРОВАННО из таблицы (в отчёте недопустимы «примерные» цифры LLM).
"""
from __future__ import annotations

import re
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path
from typing import Any, Iterable

from sqlalchemy.orm import Session

from config import settings
from data.my_documents import MyDocuments
from data.users import User
from utils.logger import logger

# Запрос «отчёт по ДПО» в чате.
DPO_REQUEST_RE = re.compile(
    r"(отч[её]т\w*[^.\n]{0,50}(\bдпо\b|повышени\w*\s+квалификац))"
    r"|((\bдпо\b|повышени\w*\s+квалификац)[^.\n]{0,50}отч[её]т)",
    re.IGNORECASE,
)

_HEADER_MARKERS = ("физическое лицо", "категория должности")
_SHORT_HOURS = 16  # порог «краткосрочной» программы

# Категории должностей → группы образца отчёта (считаются по УНИКАЛЬНЫМ людям).
_CATEGORY_GROUPS: tuple[tuple[str, tuple[str, ...], str], ...] = (
    ("ППС", ("ППС",), "профессорско-преподавательский состав (ППС)"),
    ("ПС и ПР", ("ПС", "ПР"), "педагогические работники СПО и СОО (ПС и ПР)"),
    ("НР", ("НР", "НТР"), "научные и научно-технические работники (НР, НТР)"),
    ("АУП", ("АУП", "АУПН"), "административно-управленческий персонал (АУП, АУПН)"),
    ("УВП", ("УВП", "ПУП"), "учебно-вспомогательный персонал (УВП, ПУП)"),
    ("ИТР", ("ИТР", "ПРОП"), "инженерно-технические работники (ИТР, ПрОП)"),
    ("АХП", ("АХП", "АХПН"), "административно-хозяйственный персонал (АХП, АХПН)"),
)

# Виды образования (по записям с часами >= 16, как в образце).
_KIND_LABELS = (
    ("Профессиональная переподготовка", "по программам профессиональной переподготовки"),
    ("Профессиональное обучение", "по программам профессионального обучения"),
    ("Повышение квалификации", "по программам повышения квалификации"),
)

# Обязательные программы (перечень образца) → ключи поиска в «Теме».
_MANDATORY = (
    ("Оказание первой помощи", re.compile(r"перв\w*\s+помощ", re.I)),
    ("Противодействие коррупции", re.compile(r"коррупц", re.I)),
    ("Обучение инвалидов и лиц с ОВЗ", re.compile(r"инвалид|ограниченными возможностями|\bОВЗ\b", re.I)),
    ("Охрана труда", re.compile(r"охран\w*\s+труда", re.I)),
    ("Контрактная система в сфере закупок", re.compile(r"закупок|контрактн\w*\s+систем", re.I)),
    ("Гражданская оборона и защита от ЧС", re.compile(r"гражданск\w*\s+оборон", re.I)),
    ("Информационно-образовательная среда", re.compile(r"информационно-образовательн", re.I)),
)

_YEAR_RE = re.compile(r"\b(20\d{2})\b")


def analyze_dpo_rows(rows: Iterable[tuple]) -> dict[str, Any]:
    """Агрегация выгрузки. rows — строки листа (кортежи значений), включая шапку."""
    header_idx: dict[str, int] = {}
    records: list[dict[str, Any]] = []
    current_fio: str | None = None
    param_year: int | None = None

    for row in rows:
        cells = ["" if c is None else str(c).strip() for c in row]
        if not any(cells):
            continue
        low = [c.lower() for c in cells]
        if not header_idx:
            # Параметры до шапки: «Период завершения: 31.12.2023» → год отчёта
            joined = " ".join(cells)
            if "период" in joined.lower():
                m = _YEAR_RE.search(joined)
                if m:
                    param_year = int(m.group(1))
            if any(mk in low for mk in _HEADER_MARKERS):
                header_idx = {name: i for i, name in enumerate(low) if name}
            continue

        def col(name: str) -> str:
            i = header_idx.get(name)
            return cells[i] if i is not None and i < len(cells) else ""

        # forward-fill ФИО (1С показывает его только в первой строке блока)
        fio = col("физическое лицо")
        if fio:
            current_fio = fio
        if not (col("вид образования") or col("тема") or col("категория должности")):
            continue  # служебная строка

        hours_raw = col("часов").replace(",", ".")
        try:
            hours = float(hours_raw) if hours_raw else None
        except ValueError:
            hours = None

        records.append({
            "fio": current_fio or f"(без ФИО #{len(records)})",
            "category": col("категория должности").upper().replace("РОП", "РОП"),
            "kind": col("вид образования"),
            "group": col("группа"),
            "theme": col("тема"),
            "org": col("учереждение") or col("учреждение"),
            "hours": hours,
            "issued": col("дата выдачи"),
        })

    if not records:
        raise ValueError(
            "Не удалось разобрать выгрузку: не найдена шапка таблицы "
            "(«Физическое лицо», «Категория должности», …)"
        )

    people: dict[str, list[dict]] = defaultdict(list)
    for r in records:
        people[r["fio"]].append(r)

    long_recs = [r for r in records if r["hours"] is not None and r["hours"] >= _SHORT_HOURS]
    short_recs = [r for r in records if r["hours"] is not None and r["hours"] < _SHORT_HOURS]

    kind_counts = Counter(r["kind"] for r in long_recs)
    kinds = [(label, kind_counts.get(kind, 0)) for kind, label in _KIND_LABELS]

    # Категории — по уникальным людям (категория человека = самая частая у его записей)
    person_cat: dict[str, str] = {}
    for fio, recs in people.items():
        cats = Counter(r["category"] for r in recs if r["category"])
        person_cat[fio] = cats.most_common(1)[0][0] if cats else ""
    cat_people: list[tuple[str, int]] = []
    for _key, codes, label in _CATEGORY_GROUPS:
        n = sum(1 for c in person_cat.values() if c in codes)
        cat_people.append((label, n))
    known_codes = {c for _k, codes, _l in _CATEGORY_GROUPS for c in codes}
    other_people = sum(1 for c in person_cat.values() if c and c not in known_codes)

    # Формы обучения — по людям (человек может попасть в несколько форм)
    def _people_in(group_re: re.Pattern) -> int:
        return sum(
            1 for recs in people.values()
            if any(group_re.search(r["group"] or "") for r in recs)
        )

    forms = {
        "internal": _people_in(re.compile(r"внутривуз", re.I)),
        "external": _people_in(re.compile(r"иные", re.I)),
        "internship": _people_in(re.compile(r"стажиров", re.I)),
    }

    mandatory = []
    for label, rx in _MANDATORY:
        n_people = sum(1 for recs in people.values() if any(rx.search(r["theme"] or "") for r in recs))
        if n_people:
            mandatory.append((label, n_people))

    theme_people: dict[str, set] = defaultdict(set)
    for r in records:
        if r["theme"]:
            theme_people[r["theme"]].add(r["fio"])
    top_themes = sorted(theme_people.items(), key=lambda kv: len(kv[1]), reverse=True)[:10]

    year = param_year
    if not year:
        years = Counter(
            int(m.group(1)) for r in records
            if (m := _YEAR_RE.search(r["issued"] or ""))
        )
        year = years.most_common(1)[0][0] if years else datetime.now().year

    multi = sum(1 for recs in people.values() if len(recs) >= 2)
    return {
        "year": year,
        "total_records": len(records),
        "total_people": len(people),
        "multi_program_people": multi,
        "total_programs": len(theme_people),
        "long_events": len(long_recs),
        "short_events": len(short_recs),
        "kinds": kinds,                      # [(подпись, число мероприятий ≥16ч)]
        "categories": cat_people,            # [(подпись, человек)]
        "categories_other": other_people,
        "forms": forms,
        "mandatory": mandatory,              # [(программа, человек)]
        "top_themes": [(t, len(p)) for t, p in top_themes],
    }


def analyze_dpo_xlsx(path: str | Path) -> dict[str, Any]:
    from openpyxl import load_workbook

    wb = load_workbook(str(path), read_only=True, data_only=True)
    try:
        ws = wb.worksheets[0]
        return analyze_dpo_rows(ws.iter_rows(values_only=True))
    finally:
        wb.close()


def build_report_text(stats: dict[str, Any]) -> str:
    """Текст отчёта по структуре образца «ДПО за 2023». Все числа — из таблицы."""
    y = stats["year"]
    p: list[str] = []
    p.append(f"Отчет по ДПО за {y} год")
    p.append(
        f"1. Дополнительное профессиональное образование работников университета в {y} году "
        "осуществлялось в соответствии с Порядком организации дополнительного профессионального "
        "образования и внутриорганизационного обучения работников университета, разделом 8 "
        "Коллективного договора ТИУ, разделом 9 Трудового кодекса Российской Федерации и "
        "Федеральным законом от 29.12.2012 № 273-ФЗ «Об образовании в Российской Федерации»."
    )
    p.append(
        f"В {y} году было организовано дополнительное профессиональное образование "
        f"{stats['total_people']} работников по {stats['total_programs']} программам обучения "
        f"(при этом {stats['multi_program_people']} человек — по 2 и более программам)."
    )
    kinds_lines = "\n".join(f"{n} — {label};" for label, n in stats["kinds"])
    p.append(
        f"Всего проведено {stats['long_events']} обучающих мероприятий по программам "
        f"дополнительного профессионального образования (от 16 часов), в том числе:\n{kinds_lines}"
    )
    if stats["short_events"]:
        p.append(f"Проведено {stats['short_events']} краткосрочных программ (до 16 часов).")

    cat_lines = "\n".join(
        f"{label} – {n} человек;" for label, n in stats["categories"] if n
    )
    if stats.get("categories_other"):
        cat_lines += f"\nиные категории – {stats['categories_other']} человек;"
    p.append(f"По категориям должностей обучение проходили:\n{cat_lines}")

    f = stats["forms"]
    p.append(
        "При организации повышения квалификации использовались разные формы обучения:\n"
        f"внутривузовское повышение квалификации (по программам ИДДО) – {f['internal']} человек;\n"
        f"курсы повышения квалификации, программы профессиональной переподготовки в других "
        f"образовательных организациях – {f['external']} человек;\n"
        f"стажировки – {f['internship']} человек."
    )
    if stats["mandatory"]:
        mand = "\n".join(f"{label} – {n} человек;" for label, n in stats["mandatory"])
        p.append(
            f"В {y} году повышение квалификации работников осуществлялось по программам, "
            "входящим в перечень обязательного обучения в соответствии с действующим "
            f"законодательством, в том числе:\n{mand}"
        )
    if stats["top_themes"]:
        top = "\n".join(f"«{t[:120]}» – {n} человек;" for t, n in stats["top_themes"])
        p.append(f"Наиболее массовые программы обучения {y} года:\n{top}")
    p.append(
        "[Разделы о программах по приоритетным направлениям развития университета, "
        "бесплатных онлайн-программах и выполнении плана ВОО заполняются вручную — "
        "этих данных нет в выгрузке 1С.]"
    )
    return "\n\n".join(p)


def render_dpo_docx(stats: dict[str, Any], body_text: str) -> Path:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    paras = re.split(r"\n\s*\n+", body_text)
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.add_run(paras[0])
    run.bold = True
    run.font.size = Pt(14)
    for para in paras[1:]:
        lines = para.split("\n")
        p0 = doc.add_paragraph(lines[0])
        p0.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for extra in lines[1:]:
            li = doc.add_paragraph(extra)
            li.paragraph_format.left_indent = Pt(24)

    settings.docs_generated.mkdir(parents=True, exist_ok=True)
    out = settings.docs_generated / f"dpo_report_{stats['year']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(str(out))
    return out


def create_dpo_report(db: Session, user: User, xlsx_path: str | Path) -> tuple[MyDocuments, str, dict]:
    """Полный цикл: xlsx → агрегаты → текст → docx → «Мои документы»."""
    stats = analyze_dpo_xlsx(xlsx_path)
    text = build_report_text(stats)
    path = render_dpo_docx(stats, text)
    rec = MyDocuments(
        user_id=user.id,
        title=f"Отчёт по ДПО за {stats['year']} год",
        template_key="dpo_report",
        file_path=str(path),
        progress=100,
        status="ready",
        fields={"year": stats["year"], "total_people": stats["total_people"],
                "long_events": stats["long_events"]},
    )
    db.add(rec)
    db.commit()
    db.refresh(rec)
    logger.info("[DPO] отчёт создан: doc={} ({} записей, {} человек)",
                rec.id, stats["total_records"], stats["total_people"])
    return rec, text, stats
