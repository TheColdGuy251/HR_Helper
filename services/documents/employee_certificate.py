"""Б3: справка на работника из выгрузки 1С:ЗиК → читабельный docx.

Выгрузка — xls «СПРАВКА НА СОТРУДНИКА»: пары «поле → значение» (значения
многострочные, «дата,текст,место»). Правки по брифу УРП:
- «Повышение квалификации» — только за последние 3 года;
- «Работа по окончании ВУЗа» — по ДОЛЖНОСТЯМ: подряд идущие приказы по той же
  должности схлопываются, остаётся первая дата («01.01.2012 – Ассистент, …»);
  служебные хвосты приказов («0,5 Перемещение Основное место работы») убираются.

Всё детерминированно, без LLM.
"""

from __future__ import annotations

import re
from datetime import date, datetime, timedelta
from pathlib import Path

from sqlalchemy.orm import Session

from config import settings
from data.my_documents import MyDocuments
from data.users import User
from utils.logger import logger

# Триггер чат-команды: «сделай справку на работника из выгрузки»
CERTIFICATE_EMP_REQUEST_RE = re.compile(
    r"справк\w*\s+на\s+(?:работник|сотрудник)|читабельн\w+\s+справк|преобразу\w+\s+справк",
    re.IGNORECASE,
)

_DATE_RE = re.compile(r"^(\d{2})\.(\d{2})\.(\d{4})\b")
# Служебный хвост записи о работе: «0,5 Перемещение Основное место работы» и т.п.
_WORK_TAIL_RE = re.compile(
    r"[\s,]*\d+(?:[.,]\d+)?\s*(?:Перемещение|Прием|Приём|Увольнение)?"
    r"\s*(?:Основное место работы|Внутреннее совместительство|Внешнее совместительство|Совместительство)?\s*$",
    re.IGNORECASE,
)
_PK_KEEP_YEARS = 3

# Порядок и подписи полей в итоговой справке
_FIELD_ORDER = [
    "ФИО", "Дата рождения", "Занимаемая должность", "Структурное подразделение",
    "Ученое звание", "Ученая степень", "Преподаваемые дисциплины", "Телефоны",
    "Образование", "Профессиональная переподготовка", "Повышение квалификации",
    "Работа по окончании ВУЗа", "Общий стаж", "Общий научно-педагогический стаж",
    "Стаж работы в ТИУ", "Поощрения и награды",
]
# Многострочные поля: каждая строка значения — отдельная запись-буллет
_LIST_FIELDS = {
    "Образование", "Повышение квалификации", "Профессиональная переподготовка",
    "Работа по окончании ВУЗа", "Поощрения и награды", "Преподаваемые дисциплины",
}


def _norm_label(label: str) -> str:
    """«Образование ⏎(ВУЗ, год…)» → «Образование» — сводим к каноническому имени."""
    head = re.split(r"[(\n]", label or "")[0].strip().rstrip(" :")
    for canon in _FIELD_ORDER:
        if head.lower().startswith(canon.lower()):
            return canon
    return head


def _rec_date(line: str) -> date | None:
    m = _DATE_RE.match(line.strip())
    if not m:
        return None
    try:
        return date(int(m.group(3)), int(m.group(2)), int(m.group(1)))
    except ValueError:
        return None


def parse_certificate_xls(path: str | Path) -> dict[str, str]:
    """Пары «поле → значение» из выгрузки (первый лист; .xls конвертируется)."""
    from openpyxl import load_workbook

    p = Path(path)
    if p.suffix.lower() == ".xls":
        from services.parsers.office_convert import convert_to_modern

        p = convert_to_modern(p)
    ws = load_workbook(str(p), data_only=True).active
    fields: dict[str, str] = {}
    for row in ws.iter_rows(values_only=True):
        label = str(row[0]).strip() if row and row[0] is not None else ""
        value = str(row[1]).strip() if row and len(row) > 1 and row[1] is not None else ""
        if not label or "справка" in label.lower():
            continue
        fields[_norm_label(label)] = value
    if not fields.get("ФИО"):
        raise ValueError("Файл не похож на выгрузку «Справка на сотрудника» из 1С:ЗиК")
    return fields


def _filter_pk(lines: list[str], today: date | None = None) -> list[str]:
    """Повышение квалификации: только записи за последние 3 года."""
    cutoff = (today or date.today()) - timedelta(days=_PK_KEEP_YEARS * 365)
    out = []
    for ln in lines:
        d = _rec_date(ln)
        if d is None or d >= cutoff:
            out.append(ln)
    return out


def _dedup_work(lines: list[str]) -> list[str]:
    """Работа по окончании ВУЗа → по должностям: подряд идущие записи одной
    должности схлопываются, остаётся ПЕРВАЯ дата. Хвосты приказов убираются."""
    recs: list[tuple[str, str]] = []  # (дата-строка, «Должность, Подразделение»)
    for ln in lines:
        ln = ln.strip()
        m = _DATE_RE.match(ln)
        if not m:
            if recs:  # перенос без даты — продолжение предыдущей записи
                recs[-1] = (recs[-1][0], f"{recs[-1][1]} {ln}".strip())
            continue
        rest = ln[m.end():].lstrip(" ,;–-")
        rest = _WORK_TAIL_RE.sub("", rest).strip(" ,;")
        rest = re.sub(r"\s+", " ", rest)
        recs.append((m.group(0), rest))

    out: list[str] = []
    prev_key = None
    for d, rest in recs:
        key = re.sub(r"[^а-яёa-z0-9]+", "", rest.lower())
        if key and key == prev_key:
            continue  # тот же состав должности/подразделения — пропускаем приказ
        prev_key = key
        out.append(f"{d} – {rest}")
    return out


def build_certificate_fields(raw: dict[str, str], today: date | None = None) -> dict[str, list[str] | str]:
    """Применяет правила брифа и раскладывает значения по виду (строка/список)."""
    out: dict[str, list[str] | str] = {}
    for name in _FIELD_ORDER:
        val = (raw.get(name) or "").strip()
        if not val:
            continue
        if name in _LIST_FIELDS:
            lines = [ln.strip().rstrip(";") for ln in val.splitlines() if ln.strip()]
            if name == "Повышение квалификации":
                lines = _filter_pk(lines, today)
            elif name == "Работа по окончании ВУЗа":
                lines = _dedup_work(lines)
            else:
                lines = [re.sub(r"^(\d{2}\.\d{2}\.\d{4})[\s,]+", r"\1 – ", ln) for ln in lines]
            if lines:
                out[name] = lines
        else:
            out[name] = val
    return out


def render_certificate_docx(fields: dict[str, list[str] | str]) -> Path:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    head = doc.add_paragraph()
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = head.add_run("СПРАВКА НА СОТРУДНИКА")
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph()

    for name in _FIELD_ORDER:
        if name not in fields:
            continue
        val = fields[name]
        p = doc.add_paragraph()
        p.add_run(f"{name}: ").bold = True
        if isinstance(val, str):
            p.add_run(val)
        else:
            for ln in val:
                li = doc.add_paragraph(ln)
                li.paragraph_format.left_indent = Pt(20)
                li.paragraph_format.space_after = Pt(2)

    settings.docs_generated.mkdir(parents=True, exist_ok=True)
    out = settings.docs_generated / f"certificate_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(str(out))
    return out


def create_certificate(db: Session, user: User, xls_path: str | Path) -> tuple[MyDocuments, dict]:
    raw = parse_certificate_xls(xls_path)
    fields = build_certificate_fields(raw)
    path = render_certificate_docx(fields)
    fio = str(raw.get("ФИО") or "").strip()
    rec = MyDocuments(
        user_id=user.id,
        title=f"Справка на сотрудника: {fio}" if fio and "Х" not in fio[:3] else "Справка на сотрудника",
        template_key="employee_certificate",
        file_path=str(path),
        progress=100,
        status="ready",
        fields={"fio": fio},
    )
    db.add(rec)
    db.commit()
    db.refresh(rec)
    logger.info("[CERT] справка создана: doc={}", rec.id)
    return rec, fields
