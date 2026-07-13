"""Б5: объявление на конкурс ППС из выгрузок 1С:ЗиК «Форма 2».

Кадры → Аттестации → Конкурсные отборы и выборы → печать → форма 2: по одному
файлу на должность («Конкурс: на должность - Доцент, …»). В файле люди
сгруппированы строками-заголовками институт → кафедра; по каждому — учёная
степень и специальность.

Объявление собирается по образцу УРП: секция «ВЫБОРЫ ЗАВЕДУЮЩИХ КАФЕДРАМИ»,
затем «КОНКУРС НА ЗАМЕЩЕНИЕ ДОЛЖНОСТЕЙ» по типам (профессора → доценты →
старшие преподаватели → преподаватели → ассистенты); внутри — кафедры со
скобкой требований, выведенных из данных выгрузки (специальность + степень
переизбираемых работников). Скобки — черновик для правки кадровиком в Word.

Всё детерминированно, без LLM.
"""

from __future__ import annotations

import re
from datetime import date, datetime
from pathlib import Path

from sqlalchemy.orm import Session

from config import settings
from data.my_documents import MyDocuments
from data.users import User
from utils.logger import logger

# Триггер чат-команды: «сделай объявление о конкурсе ППС», «выборы завкафедрами»
PPS_REQUEST_RE = re.compile(
    r"объявлени\w*[^.]{0,50}(?:конкурс|ппс|выбор)|конкурс[^.]{0,20}ппс|выбор\w*\s+завед",
    re.IGNORECASE,
)

# Порядок секций объявления; ключ — по подстроке в названии должности
_SECTION_ORDER = [
    ("заведующ", "ВЫБОРЫ\nЗАВЕДУЮЩИХ КАФЕДРАМИ:"),
    ("профессор", "ПРОФЕССОРОВ КАФЕДР:"),
    ("доцент", "ДОЦЕНТОВ КАФЕДР:"),
    ("старший преподаватель", "СТАРШИХ ПРЕПОДАВАТЕЛЕЙ КАФЕДР:"),
    ("преподаватель", "ПРЕПОДАВАТЕЛЕЙ КАФЕДР:"),
    ("ассистент", "АССИСТЕНТОВ КАФЕДР:"),
]

_TITLE_RE = re.compile(r"на должность\s*[-–—]\s*(.+)", re.IGNORECASE)
_DEGREE_RE = re.compile(r"(доктор|кандидат)\s+[а-яё-]+\s+наук", re.IGNORECASE)


def _section_key(position: str) -> str:
    low = (position or "").lower()
    for key, _ in _SECTION_ORDER:
        if key in low:
            return key
    return "прочие"


def parse_form2_xlsx(path: str | Path) -> dict:
    """Один файл «Форма 2» → {position, rows:[{fio, institute, department,
    degree, specialties}]}. Группировка институт/кафедра — по строкам-заголовкам."""
    from openpyxl import load_workbook

    p = Path(path)
    if p.suffix.lower() == ".xls":
        from services.parsers.office_convert import convert_to_modern

        p = convert_to_modern(p)
    ws = load_workbook(str(p), data_only=True).active

    position = None
    header_row = None
    cols: dict[str, int] = {}
    for ri, row in enumerate(ws.iter_rows(max_row=25, values_only=True), 1):
        joined = " ".join(str(v) for v in row if v is not None)
        if position is None:
            m = _TITLE_RE.search(joined)
            if m:
                position = re.sub(r"\s+", " ", m.group(1)).strip()
        low = [str(v).strip().lower() if v is not None else "" for v in row]
        if any("фамилия" in c for c in low):
            header_row = ri
            for ci, c in enumerate(low):
                if "фамилия" in c:
                    cols["fio"] = ci
                elif "ученая степень" in c or "учёная степень" in c:
                    cols["degree"] = ci
                elif "специальность" in c:
                    cols["spec"] = ci
                elif c == "должность":
                    cols["position"] = ci
            break
    if header_row is None or "fio" not in cols:
        raise ValueError(f"{Path(path).name}: не похоже на «Форму 2» (нет шапки с ФИО)")

    rows: list[dict] = []
    institute = department = ""
    for row in ws.iter_rows(min_row=header_row + 1, values_only=True):
        vals = [v for v in row if v is not None and str(v).strip()]
        if not vals:
            continue
        # Строка-заголовок группы (институт/кафедра) — объединённая ячейка:
        # в строке ровно одно непустое значение, в какой бы колонке оно ни лежало.
        if len(vals) == 1:
            text = re.sub(r"\s+", " ", str(vals[0])).strip()
            if re.match(r"(базовая\s+)?кафедра\b", text, re.IGNORECASE):
                department = text
            else:
                institute = text
                department = ""
            continue
        fio = row[cols["fio"]] if cols["fio"] < len(row) else None
        if fio is None or not str(fio).strip():
            continue
        degree = str(row[cols["degree"]]).strip() if cols.get("degree") is not None and cols["degree"] < len(row) and row[cols["degree"]] else ""
        spec = str(row[cols["spec"]]).strip() if cols.get("spec") is not None and cols["spec"] < len(row) and row[cols["spec"]] else ""
        rows.append({
            "fio": re.sub(r"\s+", " ", str(fio)).strip(),
            "institute": institute,
            "department": department or institute,
            "degree": degree,
            "specialties": spec,
        })
    if not rows:
        raise ValueError(f"{Path(path).name}: в «Форме 2» не нашлось строк с работниками")
    return {"position": position or "", "rows": rows}


def _dept_display(department: str) -> str:
    """«Кафедра интеллектуальных систем и технологий» → «интеллектуальных систем и технологий»."""
    d = re.sub(r"^(базовая\s+)?кафедра\s+", "", department.strip(), flags=re.IGNORECASE)
    return d[:1].lower() + d[1:] if d else department


def _person_profile(row: dict) -> str:
    """Черновик требований по одному работнику: образование + степень."""
    bits: list[str] = []
    specs = [s.strip(" -–") for s in re.split(r"[\n;]", row["specialties"]) if s.strip()]
    if specs:
        # «Программное обеспечение…-Инженер» → специальность до «-квалификация»
        names = []
        for s in specs[:2]:
            name = re.split(r"\s*-\s*(?=[А-ЯЁ][а-яё]+$)", s)[0].strip()
            if name and name.lower() not in (n.lower() for n in names):
                names.append(name)
        if names:
            quoted = ", ".join(f"«{n}»" for n in names)
            bits.append(
                f"образование высшее по специальности {quoted}" if len(names) == 1
                else f"образование высшее по специальностям: {quoted}"
            )
    else:
        bits.append("образование высшее")
    m = _DEGREE_RE.search(row["degree"] or "")
    if m:
        bits.append(m.group(0).lower())
    elif row["degree"]:
        bits.append(row["degree"].lower())
    return ", ".join(bits)


def build_announcement(form2_list: list[dict], announce_date: date | None = None) -> dict:
    """Собирает секции объявления из разобранных «форм 2»."""
    announce_date = announce_date or date.today()
    sections: dict[str, dict[str, list[str]]] = {}
    people = 0
    for f2 in form2_list:
        key = _section_key(f2["position"])
        depts = sections.setdefault(key, {})
        for row in f2["rows"]:
            people += 1
            disp = _dept_display(row["department"])
            profile = _person_profile(row)
            if key == "заведующ":
                profile = "наличие ученой степени и ученого звания"
            lst = depts.setdefault(disp, [])
            if profile and profile not in lst:
                lst.append(profile)

    ordered: list[tuple[str, list[tuple[str, str]]]] = []
    for key, header in _SECTION_ORDER:
        if key not in sections:
            continue
        lines = [
            (dept, "; ".join(profiles))
            for dept, profiles in sorted(sections[key].items())
        ]
        ordered.append((header, lines))
    return {
        "date": announce_date.strftime("%d.%m.%Y"),
        "sections": ordered,
        "positions": len(form2_list),
        "departments": len({d for s in sections.values() for d in s}),
        "people": people,
    }


_FOOTER_PARAGRAPHS = [
    "Претенденты для участия в выборах на должности заведующих кафедрами предоставляют "
    "документы в соответствии с пунктами 3.5 и 3.6 Порядка выборов на должность заведующего "
    "кафедрой ТИУ от 29.06.2020 года (с изменениями от 10.08.2022 года), размещенного на сайте ТИУ.",
    "Заявления и необходимые документы для участия в выборах направлять по адресу "
    "г. Тюмень, ул. Володарского, 38, каб. 106.",
    "Претенденты для участия в конкурсе на должности педагогических работников, относящихся "
    "к профессорско-преподавательскому составу (ППС), предоставляют документы в соответствии "
    "с Положением о порядке замещения должностей педагогических работников, относящихся к "
    "профессорско-преподавательскому составу (утверждено приказом Министерства науки и высшего "
    "образования РФ от 04.12.2023 г. № 1138), и Порядком замещения должностей педагогических "
    "работников ТИУ, относящихся к профессорско-преподавательскому составу, размещенными на сайте ТИУ.",
    "С претендентами, прошедшими конкурс на замещение должностей ППС, заключается трудовой "
    "договор/дополнительное соглашение к трудовому договору на срок не менее трёх лет и не "
    "более пяти лет.",
    "Заявления и необходимые документы для участия в конкурсе направлять на согласование в "
    "системе 1С: Документооборот (претендентам, не являющимся работниками ТИУ, — направлять "
    "по адресу г. Тюмень, ул. Володарского, 38, каб. 106).",
    "Срок подачи заявлений об участии в выборах и конкурсе – один месяц со дня опубликования "
    "объявления.",
    "тел. для справок: 28-35-60, вн. 11-35",
]


def render_announcement_docx(data: dict) -> Path:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    for line in (
        "ФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ БЮДЖЕТНОЕ ОБРАЗОВАТЕЛЬНОЕ УЧРЕЖДЕНИЕ ВЫСШЕГО ОБРАЗОВАНИЯ",
        "«ТЮМЕНСКИЙ ИНДУСТРИАЛЬНЫЙ УНИВЕРСИТЕТ»",
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(line).bold = True

    row = doc.add_paragraph()
    row.paragraph_format.tab_stops.add_tab_stop(Pt(460))
    row.add_run(f"г. Тюмень\t{data['date']}")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ОБЪЯВЛЯЕТ")
    r.bold = True
    r.font.size = Pt(14)

    first_konkurs = True
    for header, lines in data["sections"]:
        if not header.startswith("ВЫБОРЫ") and first_konkurs:
            ph = doc.add_paragraph()
            ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ph.add_run("КОНКУРС НА ЗАМЕЩЕНИЕ ДОЛЖНОСТЕЙ").bold = True
            first_konkurs = False
        for hline in header.split("\n"):
            ph = doc.add_paragraph()
            ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ph.add_run(hline).bold = True
        for dept, req in lines:
            li = doc.add_paragraph(f"{dept} ({req});")
            li.paragraph_format.left_indent = Pt(18)
            li.paragraph_format.space_after = Pt(2)
        doc.add_paragraph()

    for text in _FOOTER_PARAGRAPHS:
        doc.add_paragraph(text)

    settings.docs_generated.mkdir(parents=True, exist_ok=True)
    out = settings.docs_generated / f"pps_announcement_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(str(out))
    return out


def create_announcement(
    db: Session, user: User, paths: list[str | Path]
) -> tuple[MyDocuments, dict]:
    """Полный цикл: файлы «Форма 2» (по одному на должность) → word-объявление."""
    form2_list = [parse_form2_xlsx(p) for p in paths]
    data = build_announcement(form2_list)
    path = render_announcement_docx(data)
    rec = MyDocuments(
        user_id=user.id,
        title=f"Объявление о конкурсе ППС от {data['date']}",
        template_key="pps_announcement",
        file_path=str(path),
        progress=100,
        status="ready",
        fields={"positions": data["positions"], "departments": data["departments"]},
    )
    db.add(rec)
    db.commit()
    db.refresh(rec)
    logger.info("[PPS] объявление создано: doc={} ({} должностей, {} кафедр)",
                rec.id, data["positions"], data["departments"])
    return rec, data
