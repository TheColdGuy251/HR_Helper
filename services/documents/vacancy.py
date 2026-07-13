"""Б6: текст вакансии для job-сайтов из должностной инструкции.

Руководитель прикладывает ДИ; раздел 2 «Основные должностные обязанности
(трудовая функция)» переводится LLM в удобную форму объявления: обязанности
живым языком, требования (аккуратно выведенные из обязанностей), стандартный
блок условий ТИУ. Зарплата и конкретные условия не выдумываются — их
заполняет специалист УРП перед публикацией.
"""

from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path
from typing import Any

from sqlalchemy.orm import Session

from config import settings
from data.my_documents import MyDocuments
from data.users import User
from services.llm import get_llm
from utils.logger import logger

# Триггер чат-команды: «сделай вакансию по этой инструкции», «текст для hh»
VACANCY_REQUEST_RE = re.compile(
    r"(ваканси|объявлени\w+\s+(?:на|о)\s+(?:работу|должность)|"
    r"текст\s+для\s+(?:hh|джоб|job|хедхантер))",
    re.IGNORECASE,
)

# Раздел 2 ДИ: «2 ОСНОВНЫЕ ДОЛЖНОСТНЫЕ ОБЯЗАННОСТИ (ТРУДОВАЯ ФУНКЦИЯ)» /
# «2. Должностные обязанности» / «II. Трудовые функции»
_SECTION2_RE = re.compile(
    r"(?:^|\n)\s*(?:2|II)[\s.)-]+[^\n]*"
    r"(?:обязанност|трудов\w+\s+функци)[^\n]*\n",
    re.IGNORECASE,
)
_NEXT_SECTION_RE = re.compile(
    r"(?:^|\n)\s*(?:3|III)[\s.)-]+[^\n]{0,80}(?:прав|ответственност|взаимоотношени|связи)",
    re.IGNORECASE,
)
_POSITION_RE = re.compile(
    r"должностн\w+\s+инструкци\w+\s*\n?\s*([^\n]{3,90})", re.IGNORECASE
)

SYSTEM_PROMPT_VACANCY = (
    "Ты — HR-специалист Тюменского индустриального университета (ТИУ), готовишь текст "
    "вакансии для публикации на job-сайтах (hh.ru и т.п.).\n"
    "Тебе дают должность и раздел «Должностные обязанности» из должностной инструкции.\n"
    "Правила:\n"
    "1. Пиши живым, уважительным языком без канцелярита; убирай номера пунктов (2.1, 2.2), "
    "дублирование и общие формальности («соблюдает правила внутреннего распорядка», "
    "«выполняет иные поручения» — не включай).\n"
    "2. Структура строго:\n"
    "Чем предстоит заниматься:\n- 6–10 ёмких пунктов по содержанию обязанностей\n\n"
    "Что мы ожидаем:\n- 4–6 требований, АККУРАТНО выведенных из обязанностей "
    "(навыки, инструменты, качества); ничего не выдумывай сверх текста\n\n"
    "Условия:\n- работа в крупнейшем техническом университете Тюменской области\n"
    "- официальное трудоустройство по ТК РФ, стабильные выплаты\n"
    "- социальная программа университета (льготы, поддержка сотрудников)\n"
    "3. НЕ указывай зарплату, график и адрес — их добавит специалист УРП.\n"
    "4. Не используй markdown-заголовки (#), только строки с двоеточием и дефисные списки.\n"
    "Верни только текст объявления."
)


def _to_nominative(pos: str) -> str:
    """«специалиста ректората» → «Специалист ректората»: ведущие прилагательные
    и первое существительное приводим к именительному, дополнение («ректората»)
    остаётся в родительном — так и должно быть."""
    try:
        import pymorphy3

        morph = pymorphy3.MorphAnalyzer()
        words = pos.split()
        out: list[str] = []
        head_done = False
        for w in words:
            if head_done:
                out.append(w)
                continue
            p = morph.parse(w)[0]
            if p.tag.POS in ("NOUN", "ADJF", "PRTF"):
                nom = p.inflect({"nomn"})
                out.append(nom.word if nom else w)
                if p.tag.POS == "NOUN":
                    head_done = True
            else:
                out.append(w)
                head_done = True
        pos = " ".join(out)
    except Exception:
        pass
    return pos


def extract_position(di_text: str) -> str | None:
    """Название должности из шапки «Должностная инструкция <должности>»."""
    m = _POSITION_RE.search(di_text or "")
    if not m:
        return None
    pos = m.group(1).strip(" .;:—-")
    if not pos:
        return None
    pos = _to_nominative(pos)
    return pos[:1].upper() + pos[1:]


def extract_duties_section(di_text: str) -> str | None:
    """Раздел 2 «Должностные обязанности» ДИ (до раздела 3). None — не найден."""
    text = di_text or ""
    m = _SECTION2_RE.search(text)
    if not m:
        return None
    start = m.start()
    m2 = _NEXT_SECTION_RE.search(text, m.end())
    section = text[start:m2.start() if m2 else len(text)].strip()
    return section if len(section) >= 200 else None


def generate_vacancy_text(duties: str, position: str | None) -> str:
    user_msg = (
        f"Должность: {position or 'не указана (возьми из текста раздела)'}\n\n"
        f"Раздел «Должностные обязанности» из должностной инструкции:\n{duties[:6000]}"
    )
    llm = get_llm()
    text = (llm.generate_text(SYSTEM_PROMPT_VACANCY, user_msg, max_tokens=1200, temperature=0.4) or "").strip()
    if not text:
        raise RuntimeError("LLM не вернула текст вакансии")
    return text


def render_vacancy_docx(position: str | None, body_text: str) -> Path:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    head = doc.add_paragraph()
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = head.add_run(f"Вакансия: {position}" if position else "Текст вакансии")
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph()
    for line in body_text.splitlines():
        line = line.rstrip()
        if not line:
            doc.add_paragraph()
            continue
        p = doc.add_paragraph(line)
        if line.startswith(("-", "–", "•")):
            p.paragraph_format.left_indent = Pt(18)

    note = doc.add_paragraph()
    note_run = note.add_run(
        "Зарплата, график работы и контакты добавляются специалистом УРП перед публикацией."
    )
    note_run.italic = True
    note_run.font.size = Pt(10)

    settings.docs_generated.mkdir(parents=True, exist_ok=True)
    out = settings.docs_generated / f"vacancy_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(str(out))
    return out


def create_vacancy(
    db: Session, user: User, di_text: str, position: str | None = None
) -> tuple[MyDocuments, str, dict[str, Any]]:
    """Полный цикл: текст ДИ → раздел 2 → LLM-текст вакансии → docx → «Мои документы»."""
    pos = position or extract_position(di_text)
    duties = extract_duties_section(di_text)
    meta = {"position": pos, "section_found": bool(duties)}
    if not duties:
        # Раздел не распознан — отдаём LLM весь текст ДИ (обязанности она выделит сама)
        duties = (di_text or "")[:6000]
        if len(duties.strip()) < 200:
            raise ValueError("В файле не нашлось текста должностной инструкции")
    text = generate_vacancy_text(duties, pos)
    path = render_vacancy_docx(pos, text)
    rec = MyDocuments(
        user_id=user.id,
        title=f"Вакансия: {pos}" if pos else "Текст вакансии",
        template_key="vacancy",
        file_path=str(path),
        progress=100,
        status="ready",
        fields=meta,
    )
    db.add(rec)
    db.commit()
    db.refresh(rec)
    logger.info("[VACANCY] объявление создано: doc={} (должность: {})", rec.id, pos)
    return rec, text, meta
