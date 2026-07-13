from __future__ import annotations

import re
import xml.etree.ElementTree as ET
import zipfile
from collections import defaultdict
from pathlib import Path

from docx import Document

from services.parsers.base import ParsedDocument, format_table

_TEXT_PARTS_RE = re.compile(r"word/(header\d*\.xml|footer\d*\.xml)")
_SMARTART_RE = re.compile(r"word/diagrams/data\d*\.xml")
_BIG = 10 ** 12               # «в конец» для фигур без координаты
_EDGE_MATCH_EMU = 700000      # макс. расстояние конца стрелки до блока (~0.77 см)


def _local(tag: str) -> str:
    """Локальное имя тега без namespace: '{...}p' -> 'p'."""
    return tag.rsplit("}", 1)[-1]


def _norm(s: str) -> str:
    return re.sub(r"\s+", " ", s or "").strip().lower()


def _paragraph_text(el: ET.Element) -> str:
    """Текст одного абзаца (w:p/a:p): прогоны <…:t>, без захода во вложенные абзацы."""
    parts: list[str] = []
    for child in el:
        lc = _local(child.tag)
        if lc == "p":
            continue
        if lc == "t":
            if child.text:
                parts.append(child.text)
        elif lc in ("br", "cr", "tab"):
            parts.append(" ")
        else:
            parts.append(_paragraph_text(child))
    return "".join(parts)


def _collect_paragraphs(el: ET.Element, out: list[str], is_root: bool = False) -> None:
    """Строки-абзацы из поддерева, ПРОПУСКАЯ mc:Fallback (дубль VML), w:tbl (таблицы
    разбираются отдельно) и ВЛОЖЕННЫЕ фигуры (их текст — отдельный блок)."""
    lc = _local(el.tag)
    if lc in ("Fallback", "tbl"):
        return
    if lc in ("anchor", "inline") and not is_root:
        return
    if lc == "p":
        line = _paragraph_text(el).strip()
        if line:
            out.append(line)
    for child in el:
        _collect_paragraphs(child, out, is_root=False)


def _anchor_text(anchor: ET.Element) -> str:
    """Текст фигуры-блока: абзацы её текст-бокса, склеенные через « / » (без повторов)."""
    lines: list[str] = []
    _collect_paragraphs(anchor, lines, is_root=True)
    seen: set[str] = set()
    uniq: list[str] = []
    for ln in lines:
        k = _norm(ln)
        if k and k not in seen:
            seen.add(k)
            uniq.append(ln)
    return " / ".join(uniq)


def _own_offset(anchor: ET.Element, which: str) -> int | None:
    for e in anchor.iter():
        if _local(e.tag) == which:
            for c in e:
                if _local(c.tag) == "posOffset" and c.text:
                    try:
                        return int(c.text)
                    except ValueError:
                        return None
    return None


def _anchor_geometry(anchor: ET.Element) -> tuple[int | None, int | None, int, int, bool, bool]:
    """Координата (posOffset, EMU), размер (extent) и отражение (flipH/flipV) фигуры."""
    x = _own_offset(anchor, "positionH")
    y = _own_offset(anchor, "positionV")
    cx = cy = 0
    for e in anchor.iter():
        if _local(e.tag) == "extent":
            try:
                cx, cy = int(e.get("cx") or 0), int(e.get("cy") or 0)
            except ValueError:
                pass
            break
    flip_h = flip_v = False
    for e in anchor.iter():
        if _local(e.tag) == "xfrm":
            flip_h, flip_v = e.get("flipH") == "1", e.get("flipV") == "1"
            break
    return x, y, cx, cy, flip_h, flip_v


def _anchor_shape_info(anchor: ET.Element) -> tuple[str | None, bool, bool]:
    """prstGeom-пресет + наличие стрелок на концах (tailEnd/headEnd)."""
    prst = None
    for e in anchor.iter():
        if _local(e.tag) == "prstGeom":
            prst = e.get("prst")
            break
    tail = head = False
    for e in anchor.iter():
        lc = _local(e.tag)
        if lc == "tailEnd" and (e.get("type") or "none") != "none":
            tail = True
        elif lc == "headEnd" and (e.get("type") or "none") != "none":
            head = True
    return prst, tail, head


def _point_rect_dist(px: float, py: float, box: tuple[int, int, int, int]) -> float:
    x, y, w, h = box
    dx = max(x - px, 0, px - (x + w))
    dy = max(y - py, 0, py - (y + h))
    return (dx * dx + dy * dy) ** 0.5


def _order_by_arrows(root: ET.Element) -> list[str] | None:
    """Восстанавливает ПОСЛЕДОВАТЕЛЬНОСТЬ блоков схемы ПО СТРЕЛКАМ: блоки-фигуры с
    текстом — узлы, коннекторы/линии со стрелкой — рёбра (по концам стрелки находим
    ближайшие блоки). Порядок — топологическая сортировка графа. Возвращает None,
    если стрелок/узлов недостаточно (тогда вызывающий откатывается на координаты)."""
    nodes: list[dict] = []
    conns: list[tuple[tuple[float, float], tuple[float, float], bool, bool]] = []
    for el in root.iter():
        if _local(el.tag) not in ("anchor", "inline"):
            continue
        x, y, cx, cy, flip_h, flip_v = _anchor_geometry(el)
        if x is None or y is None:
            continue
        prst, tail, head = _anchor_shape_info(el)
        p = (prst or "").lower()
        is_conn = ("connector" in p) or p == "line" or "arrow" in p
        text = _anchor_text(el)
        if text and not is_conn:
            nodes.append({
                "id": len(nodes), "box": (x, y, cx, cy),
                "cx": x + cx / 2, "cy": y + cy / 2, "text": text,
            })
        elif is_conn and (cx or cy):
            p1 = (x + (cx if flip_h else 0), y + (cy if flip_v else 0))
            p2 = (x + (0 if flip_h else cx), y + (0 if flip_v else cy))
            conns.append((p1, p2, tail, head))

    if len(nodes) < 2 or not conns:
        return None

    def _nearest(pt: tuple[float, float]) -> tuple[dict | None, float]:
        best, best_d = None, 1e18
        for n in nodes:
            d = _point_rect_dist(pt[0], pt[1], n["box"])
            if d < best_d:
                best, best_d = n, d
        return best, best_d

    edges: list[tuple[int, int]] = []
    for p1, p2, tail, head in conns:
        # Наконечник (tailEnd) стоит на КОНЦЕ (p2) → поток p1→p2; headEnd → p2→p1.
        start_pt, end_pt = (p2, p1) if (head and not tail) else (p1, p2)
        sn, sd = _nearest(start_pt)
        en, ed = _nearest(end_pt)
        if not sn or not en or sn is en or sd > _EDGE_MATCH_EMU or ed > _EDGE_MATCH_EMU:
            continue
        # Линия без наконечника — направляем сверху вниз / слева направо.
        if not tail and not head and (sn["cy"], sn["cx"]) > (en["cy"], en["cx"]):
            sn, en = en, sn
        edges.append((sn["id"], en["id"]))
    edges = list(dict.fromkeys(edges))
    if not edges:
        return None

    # Топологическая сортировка (Kahn); источники и ветвления — по расположению.
    poskey = {n["id"]: (n["cy"], n["cx"]) for n in nodes}
    indeg = {n["id"]: 0 for n in nodes}
    adj: dict[int, list[int]] = defaultdict(list)
    for a, b in edges:
        adj[a].append(b)
        indeg[b] += 1
    queue = sorted([i for i in indeg if indeg[i] == 0], key=lambda i: poskey[i])
    order: list[int] = []
    seen: set[int] = set()
    while queue:
        i = queue.pop(0)
        if i in seen:
            continue
        seen.add(i)
        order.append(i)
        for j in sorted(adj[i], key=lambda k: poskey[k]):
            indeg[j] -= 1
            if indeg[j] == 0:
                queue.append(j)
        queue.sort(key=lambda k: poskey[k])
    # Циклы/несвязанные — по расположению.
    for n in sorted(nodes, key=lambda n: poskey[n["id"]]):
        if n["id"] not in seen:
            order.append(n["id"])
            seen.add(n["id"])

    id2text = {n["id"]: n["text"] for n in nodes}
    return [id2text[i] for i in order]


def _order_by_position(root: ET.Element) -> list[str]:
    """Fallback: блоки-фигуры по расположению (сверху вниз, слева направо)."""
    positioned: list[tuple[int, int, str]] = []
    for el in root.iter():
        if _local(el.tag) not in ("anchor", "inline"):
            continue
        text = _anchor_text(el)
        if not text:
            continue
        prst, _, _ = _anchor_shape_info(el)
        p = (prst or "").lower()
        if ("connector" in p) or p == "line" or "arrow" in p:
            continue
        x, y, _cx, _cy, _fh, _fv = _anchor_geometry(el)
        positioned.append((y if y is not None else _BIG, x if x is not None else _BIG, text))
    positioned.sort(key=lambda b: (b[0], b[1]))
    return [t for _, _, t in positioned]


def _extract_shape_blocks(path: Path) -> tuple[list[str], str]:
    """Текст фигур/схем, невидимый для python-docx. Порядок блоков — ПО СТРЕЛКАМ
    (если есть), иначе по расположению. Возвращает (блоки, метод: 'arrows'|'position'|'')."""
    method = ""
    ordered: list[str] = []
    smartart: list[str] = []
    try:
        with zipfile.ZipFile(str(path)) as z:
            names = z.namelist()
            if "word/document.xml" in names:
                try:
                    root = ET.fromstring(z.read("word/document.xml"))
                    by_arrows = _order_by_arrows(root)
                    if by_arrows is not None:
                        ordered, method = by_arrows, "arrows"
                    else:
                        ordered, method = _order_by_position(root), "position"
                except ET.ParseError:
                    pass
            # Колонтитулы — плоско (позиции для потока не важны).
            for part in (n for n in names if _TEXT_PARTS_RE.fullmatch(n)):
                try:
                    _collect_paragraphs(ET.fromstring(z.read(part)), ordered, is_root=True)
                except ET.ParseError:
                    continue
            # SmartArt: порядок задан в data.xml.
            for part in (n for n in names if _SMARTART_RE.fullmatch(n)):
                try:
                    _collect_paragraphs(ET.fromstring(z.read(part)), smartart, is_root=True)
                except ET.ParseError:
                    continue
    except (zipfile.BadZipFile, KeyError, OSError):
        return [], ""

    all_blocks = ordered + smartart
    seen: set[str] = set()
    unique: list[str] = []
    for ln in all_blocks:
        k = _norm(ln)
        if k and k not in seen:
            seen.add(k)
            unique.append(ln)
    return unique, method


def parse_docx(path: str | Path) -> ParsedDocument:
    path = Path(path)
    doc = Document(str(path))

    parts: list[str] = []

    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        rows = [[(c.text or "").strip() for c in row.cells] for row in table.rows]
        parts.extend(format_table(rows))

    # Текст из фигур/текст-боксов/SmartArt (схемы процессов) как ПОСЛЕДОВАТЕЛЬНОСТЬ
    # шагов. Дедуп против тела/таблиц.
    seen = {_norm(p) for p in parts}
    blocks, method = _extract_shape_blocks(path)
    blocks = [b for b in blocks if _norm(b) not in seen]
    if blocks:
        if method == "arrows":
            header = ("Схема процесса — шаги по порядку (реконструировано по СТРЕЛКАМ "
                      "между блоками; ветвления могут быть упрощены):")
        else:
            header = ("Схема процесса — шаги по порядку (реконструировано по расположению "
                      "блоков сверху вниз и слева направо; стрелки/ветвления могут быть неточны):")
        parts.append(header)
        parts.extend(f"{i}. {b}" for i, b in enumerate(blocks, start=1))

    return ParsedDocument(
        text="\n\n".join(parts),
        title=path.stem,
        source_uri=str(path),
        source_type="local",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
